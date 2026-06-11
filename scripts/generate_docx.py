import os
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=200, right=200): # 140 dxa = ~7pt, 200 dxa = ~10pt
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tcPr = cell._element.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    
    borders = {'top': top, 'bottom': bottom, 'left': left, 'right': right}
    for border_name, border_style in borders.items():
        if border_style:
            b_el = OxmlElement(f'w:{border_name}')
            b_el.set(qn('w:val'), border_style.get('val', 'single'))
            b_el.set(qn('w:sz'), str(border_style.get('sz', 4)))
            b_el.set(qn('w:space'), '0')
            b_el.set(qn('w:color'), border_style.get('color', 'auto'))
            tcBorders.append(b_el)
        else:
            b_el = OxmlElement(f'w:{border_name}')
            b_el.set(qn('w:val'), 'none')
            tcBorders.append(b_el)
    tcPr.append(tcBorders)

def apply_table_styles(table, header_bg="0F1923", alt_row_bg="F4F7F6", border_color="CCCCCC"):
    for i, row in enumerate(table.rows):
        is_header = (i == 0)
        bg_color = header_bg if is_header else (alt_row_bg if i % 2 == 1 else "FFFFFF")
        
        for cell in row.cells:
            set_cell_background(cell, bg_color)
            set_cell_margins(cell)
            
            # Subtle borders
            top_border = {'val': 'single', 'sz': 4, 'color': border_color} if not is_header else None
            bottom_border = {'val': 'single', 'sz': 8, 'color': "00C2A8" if is_header else border_color}
            set_cell_borders(cell, top=top_border, bottom=bottom_border, left=None, right=None)

def set_run_font(run, font_name="Raleway", size_pt=11, color_rgb=(15, 25, 35), bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*color_rgb)
    run.bold = bold
    run.italic = italic

def add_heading_styled(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18 if level > 1 else 24)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    
    run = p.add_run(text)
    
    if level == 1:
        set_run_font(run, "Raleway", size_pt=18, color_rgb=(0, 194, 168), bold=True)
        # Add a subtle bottom border/underline paragraph style effect
    elif level == 2:
        set_run_font(run, "Raleway", size_pt=14, color_rgb=(15, 25, 35), bold=True)
    elif level == 3:
        set_run_font(run, "Raleway", size_pt=12, color_rgb=(56, 89, 119), bold=True)
    return p

def add_bullet_styled(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    
    if bold_prefix:
        run_prefix = p.add_run(bold_prefix)
        set_run_font(run_prefix, "Raleway", size_pt=11, color_rgb=(15, 25, 35), bold=True)
        
    run_text = p.add_run(text)
    set_run_font(run_text, "Raleway", size_pt=11, color_rgb=(15, 25, 35))
    return p

def main():
    doc = Document()
    
    # Page Setup - Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Set default paragraph format
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Raleway'
    font.size = Pt(11)
    font.color.rgb = RGBColor(15, 25, 35)

    # ------------------ TITLE PAGE / HEADER ------------------
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(40)
    title_p.paragraph_format.space_after = Pt(10)
    title_run = title_p.add_run("CARBON TATVA")
    set_run_font(title_run, "Raleway", size_pt=32, color_rgb=(0, 194, 168), bold=True)
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(30)
    sub_run = subtitle_p.add_run("Vectorless RAG Chatbot: Architecture & Process Specification")
    set_run_font(sub_run, "Raleway", size_pt=16, color_rgb=(56, 89, 119), italic=True)
    
    divider_p = doc.add_paragraph()
    divider_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider_p.paragraph_format.space_after = Pt(120)
    div_run = divider_p.add_run("_____________________________________________________")
    set_run_font(div_run, "Raleway", size_pt=12, color_rgb=(0, 194, 168))
    
    metadata_p = doc.add_paragraph()
    metadata_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    metadata_p.paragraph_format.line_spacing = 1.3
    
    m1 = metadata_p.add_run("System Architecture and Processing Documentation\n")
    set_run_font(m1, "Raleway", size_pt=11, color_rgb=(15, 25, 35), bold=True)
    m2 = metadata_p.add_run("Target Models: Mistral Large (routing) & Mistral Large (generation)\n")
    set_run_font(m2, "Raleway", size_pt=10, color_rgb=(15, 25, 35))
    m3 = metadata_p.add_run("Domain: BEE Energy Efficiency Guide Books (Electrical & Thermal Utilities)\n")
    set_run_font(m3, "Raleway", size_pt=10, color_rgb=(15, 25, 35))
    m4 = metadata_p.add_run("Deployment Environment: Vercel Python Serverless + FastAPI")
    set_run_font(m4, "Raleway", size_pt=10, color_rgb=(15, 25, 35))
    
    doc.add_page_break()
    
    # ------------------ OVERVIEW ------------------
    add_heading_styled(doc, "1. Executive Overview", level=1)
    
    p1 = doc.add_paragraph()
    p1.paragraph_format.space_after = Pt(10)
    p1.paragraph_format.line_spacing = 1.15
    run1 = p1.add_run(
        "This document outlines the design and implementation details of Carbon Tatva's Vectorless Retrieval-Augmented Generation (RAG) chatbot. "
        "The system has been engineered to deliver reliable, highly accurate, and legally grounded question-answering capabilities using the "
        "Bureau of Energy Efficiency (BEE) Guide Books for Electrical Utilities and Thermal Utilities. "
        "Unlike traditional RAG systems that rely heavily on dense vector databases and similarity metrics, this project operates on a "
        "Vectorless architecture. It replaces similarity-based lookup with a semantic Table of Contents (TOC) tree structure, parsed and navigated "
        "using large language models (LLMs). This guarantees absolute logical alignment between the retrieved text and the original books' hierarchies."
    )
    set_run_font(run1, "Raleway", size_pt=11, color_rgb=(15, 25, 35))
    
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(14)
    run2 = p2.add_run(
        "The application has been deployed on Vercel as a serverless backend with a lightweight HTML/CSS/JS frontend styled strictly in the "
        "Carbon Tatva theme using the Raleway typeface, providing responsive performance and beautiful user interaction elements."
    )
    set_run_font(run2, "Raleway", size_pt=11, color_rgb=(15, 25, 35))

    add_heading_styled(doc, "Key Architectural Differentiators", level=2)
    
    # Table: Traditional vs Vectorless
    table1 = doc.add_table(rows=6, cols=3)
    table1.autofit = False
    
    # Set widths
    widths = [Inches(1.5), Inches(2.2), Inches(2.8)]
    for row in table1.rows:
        for i, width in enumerate(widths):
            row.cells[i].width = width

    headers = ["Feature", "Traditional RAG", "Carbon Tatva Vectorless RAG"]
    for i, h in enumerate(headers):
        cell_run = table1.cell(0, i).paragraphs[0].add_run(h)
        set_run_font(cell_run, "Raleway", size_pt=11, color_rgb=(255, 255, 255), bold=True)
        table1.cell(0, i).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    row_data = [
        ["Retrieval Method", "Vector embeddings + cosine similarity", "Hierarchical TOC tree traversal via Mistral LLM routing"],
        ["Infrastructure", "Requires specialized vector databases (e.g., Pinecone, ChromaDB, PGVector)", "Zero database overhead: flat JSON structures (`tree.json`, `pages.jsonl`)"],
        ["Chunking Strategy", "Fixed-size sliding window chunks (e.g., 500 tokens with 50-token overlap)", "Semantic page-aligned boundaries matching physical TOC subsections"],
        ["Context Selection", "Top-K nearest neighbors (based on mathematical cosine closeness)", "Logical query-routing to specific Chapters/Subsections containing structured chapters"],
        ["Grounding & Quality", "Hallucination-prone: retrieved chunks might be out-of-order or contextually disjoint", "Strict structural grounding guardrails validating that facts are linked to actual pages and sections"]
    ]

    for r_idx, row_vals in enumerate(row_data):
        for c_idx, val in enumerate(row_vals):
            p = table1.cell(r_idx + 1, c_idx).paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            cell_run = p.add_run(val)
            set_run_font(cell_run, "Raleway", size_pt=10, color_rgb=(15, 25, 35), bold=(c_idx == 0))
            
    apply_table_styles(table1)
    
    doc.add_page_break()

    # ------------------ ARCHITECTURE & PROCESSES ------------------
    add_heading_styled(doc, "2. System Architecture", level=1)
    
    # 2.1 High Level
    add_heading_styled(doc, "2.1 High-Level Architecture", level=2)
    p_arch = doc.add_paragraph()
    p_arch.paragraph_format.space_after = Pt(10)
    run_arch = p_arch.add_run(
        "The system architecture is bifurcated into two primary workflows: (1) Phase 1: Offline Ingestion, where the raw PDF files are OCR-extracted, "
        "chunked, summarized, and structured into a flat-file database; and (2) Phase 2: Runtime Query Execution, where a user's question is "
        "processed through a routing model to retrieve relevant segments and generate a cited response under strict guardrails. The architecture is represented below:"
    )
    set_run_font(run_arch, "Raleway")

    # Add a block diagram text explanation since Mermaid isn't directly visual
    p_flow_desc = doc.add_paragraph()
    p_flow_desc.paragraph_format.left_indent = Inches(0.2)
    p_flow_desc.paragraph_format.space_after = Pt(10)
    run_fd = p_flow_desc.add_run(
        "Ingestion Pipeline (Offline):\n"
        "Raw PDF Book -> Render pages to high-DPI images -> Extract text via Tesseract OCR -> "
        "LLM-powered TOC Extraction -> TOC-Aware Chunking (split text by subsection boundaries) -> "
        "LLM Summarization per Subsection -> Save tree.json (navigation) & pages.jsonl (text contents)\n\n"
        "Query Pipeline (Online):\n"
        "User Question -> FastAPI Web Server -> LLM-guided TOC Tree Routing (select subsection IDs) -> "
        "Retrieve raw text from pages.jsonl -> Pre-Generation Guardrail -> Grounded Answer Generation (Mistral) -> "
        "Post-Generation Guardrail (verify citations and context grounding) -> Final Cited Answer"
    )
    set_run_font(run_fd, "Raleway", size_pt=10, color_rgb=(56, 89, 119), italic=True)

    # 2.2 Ingestion Pipeline
    add_heading_styled(doc, "2.2 Ingestion Pipeline — Detailed Flow", level=2)
    p_ing = doc.add_paragraph()
    run_ing = p_ing.add_run(
        "Offline book ingestion is performed via PyMuPDF rendering and Tesseract OCR. A structured caching layer saves OCR progress per page to avoid redundant "
        "computational steps. Once raw text is extracted, the first 22 pages of each guide are analyzed by Mistral LLM to construct a logical hierarchical TOC structure. "
        "The full text is then dynamically partitioned into discrete nodes corresponding directly to these subsections, with page ranges stored as metadata."
    )
    set_run_font(run_ing, "Raleway")
    
    add_bullet_styled(doc, "Renders each PDF page to a 300 DPI PNG image using PyMuPDF (fitz) to capture equations, text, and labels cleanly.", bold_prefix="Page Rendering: ")
    add_bullet_styled(doc, "Applies PyTesseract OCR on the page images, saving extracted text into local cache files (ocr_cache_*.json) to enable fast restarts and adjustments.", bold_prefix="OCR Extraction: ")
    add_bullet_styled(doc, "Extracts the Table of Contents structure programmatically, generating a JSON map of chapters, subchapters, and subsections.", bold_prefix="TOC Mapping: ")
    add_bullet_styled(doc, "Executes subsection chunking: segments the raw OCR text using the derived page-level boundaries of each subsection.", bold_prefix="Chunking: ")
    add_bullet_styled(doc, "Generates concise 2-3 sentence summaries for each chunk using Mistral LLM to act as routing index targets.", bold_prefix="Summarization: ")
    
    # 2.3 Runtime Execution Sequence
    add_heading_styled(doc, "2.3 Runtime Query Execution Sequence", level=2)
    p_run = doc.add_paragraph()
    run_run = p_run.add_run(
        "Upon receiving a query, the FastAPI server acts as the central coordinator, driving retrieval, guardrail validation, and response creation. "
        "The sequence of operations is as follows:"
    )
    set_run_font(run_run, "Raleway")

    # List sequence
    add_bullet_styled(doc, "The user enters a question in the Carbon Tatva chatbot interface.", bold_prefix="1. Query Receipt: ")
    add_bullet_styled(doc, "The FastAPI backend forwards the query to the Tree Retriever.", bold_prefix="2. Delegation: ")
    add_bullet_styled(doc, "The Tree Retriever loads `tree.json` and presents the hierarchical outline (titles & summaries) to the routing model (Mistral Large).", bold_prefix="3. TOC Routing: ")
    add_bullet_styled(doc, "Mistral LLM identifies and outputs the specific subsection IDs (e.g., ELECTRICAL-C3-S2) that logically hold the answer.", bold_prefix="4. Subsection Selection: ")
    add_bullet_styled(doc, "The system fetches the exact raw text of those selected subsections from `pages.jsonl`.", bold_prefix="5. Context Fetching: ")
    add_bullet_styled(doc, "The Pre-Generation Guardrail checks if any relevant contexts were found. If none, it returns an immediate out-of-bounds response.", bold_prefix="6. Pre-Gen Check: ")
    add_bullet_styled(doc, "The assembled contexts and user query are sent to Mistral LLM, configured with a system prompt that mandates strict grounding.", bold_prefix="7. Grounded Generation: ")
    add_bullet_styled(doc, "Mistral LLM returns a structured response containing: (a) the answer text, (b) citation IDs referencing the context blocks used, and (c) a grounding flag.", bold_prefix="8. Answer Receipt: ")
    add_bullet_styled(doc, "The Post-Generation Guardrail verifies: (a) the grounding flag is true, (b) all citation IDs mapped to actual retrieved sections, and (c) no out-of-bounds statements were slipped in.", bold_prefix="9. Post-Gen Guard: ")
    add_bullet_styled(doc, "The backend formats the verified citations, mapping them to the physical book page numbers, and returns the response payload.", bold_prefix="10. Response Delivery: ")

    # 2.4 Guardrails Pipeline
    add_heading_styled(doc, "2.4 Guardrails Pipeline", level=2)
    p_gr = doc.add_paragraph()
    run_gr = p_gr.add_run(
        "The chatbot implements dual-stage guardrails (Pre-Generation and Post-Generation) to completely block LLM hallucinations. "
        "If a user asks about general topics, coding, or facts not present in the BEE guides, the system intercepts the query and refuses it, "
        "ensuring that the system behaves as a strictly closed-domain assistant."
    )
    set_run_font(run_gr, "Raleway")
    
    # 2.5 TOC Tree Structure
    add_heading_styled(doc, "2.5 TOC Tree Structure", level=2)
    p_tree = doc.add_paragraph()
    run_tree = p_tree.add_run(
        "The navigation index is stored in a clean tree structure. The root node splits into the two primary books (ELECTRICAL and THERMAL). "
        "Each book node contains a list of Chapter nodes. Chapter nodes expand into Subsection leaf nodes. Each leaf node stores:\n"
        "•  Node ID (e.g., ELECTRICAL-C1-S3)\n"
        "•  Title (e.g., S3: Power Factor Improvement)\n"
        "•  Summaries (LLM-synthesized context)\n"
        "•  Page Range (e.g., pp. 17-24)\n"
        "•  Book Source Reference"
    )
    set_run_font(run_tree, "Raleway")

    doc.add_page_break()

    # ------------------ REPOSITORY & RESPONSIBILITIES ------------------
    add_heading_styled(doc, "3. Repository Architecture & Module Matrix", level=1)
    
    p_rep = doc.add_paragraph()
    run_rep = p_rep.add_run(
        "The project is structured logically to separate the ingestion scripts, runtime API layers, indexing logic, and prompt configurations. "
        "The module responsibility matrix below outlines the files and their corresponding roles in both the Ingestion and Runtime phases:"
    )
    set_run_font(run_rep, "Raleway")

    # Table: Module Responsibility Matrix
    table2 = doc.add_table(rows=11, cols=4)
    table2.autofit = False
    
    widths2 = [Inches(1.2), Inches(1.8), Inches(3.2), Inches(0.8)]
    for row in table2.rows:
        for i, width in enumerate(widths2):
            row.cells[i].width = width

    headers2 = ["Module", "File Location", "Responsibility", "Stage"]
    for i, h in enumerate(headers2):
        cell_run = table2.cell(0, i).paragraphs[0].add_run(h)
        set_run_font(cell_run, "Raleway", size_pt=10, color_rgb=(255, 255, 255), bold=True)

    matrix_data = [
        ["Interface Layer", "app/api/server.py", "FastAPI routes (/chat, /health, /) and static folder serving.", "Runtime"],
        ["Interface Layer", "app/api/static/index.html", "Single-page chat interface styled with Carbon Tatva Navy/Teal theme and Raleway typography.", "Runtime"],
        ["Extraction Layer", "app/ingest/pdf_loader.py", "PyMuPDF rendering + Tesseract OCR extraction with JSON local caching.", "Ingestion"],
        ["Extraction Layer", "app/ingest/pipeline.py", "Orchestrates extraction, TOC mapping, chunking, and index compilation.", "Ingestion"],
        ["Index Layer", "app/pageindex/indexer.py", "LLM-powered TOC parsing, structuring, and tree creation.", "Ingestion"],
        ["Index Layer", "app/pageindex/retriever.py", "Loads tree.json and delegates to routing LLM to select relevant subsection IDs.", "Runtime"],
        ["Index Layer", "app/pageindex/store.py", "Performs file disk I/O operations for tree.json and pages.jsonl.", "Both"],
        ["AI Gateway", "app/llm/mistral_client.py", "Handles Mistral HTTP API requests with exponential backoff retries for robust operation.", "Both"],
        ["Query Engine", "app/qa/engine.py", "ChatEngine orchestrator executing TOC retrieval, guardrails, and generation.", "Runtime"],
        ["Query Engine", "app/qa/guardrails.py", "Performs pre-generation and post-generation checks to guarantee grounded answers.", "Runtime"]
    ]

    for r_idx, row_vals in enumerate(matrix_data):
        for c_idx, val in enumerate(row_vals):
            p = table2.cell(r_idx + 1, c_idx).paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            cell_run = p.add_run(val)
            set_run_font(cell_run, "Raleway", size_pt=9.5, color_rgb=(15, 25, 35), bold=(c_idx == 0))
            
    apply_table_styles(table2)

    doc.add_page_break()

    # ------------------ DETAILED HOW IT WORKS ------------------
    add_heading_styled(doc, "4. Detailed Process Walkthrough", level=1)
    
    add_heading_styled(doc, "4.1 Phase 1: Ingestion Process", level=2)
    
    p_proc1 = doc.add_paragraph()
    p_proc1.paragraph_format.space_after = Pt(10)
    run_pp1 = p_proc1.add_run(
        "The offline ingestion process parses the textbooks into structured search indices. "
        "This is performed once, and the resulting pre-built index files are committed directly to the code repository. "
        "The pipeline runs as follows:"
    )
    set_run_font(run_pp1, "Raleway")

    add_bullet_styled(doc, "The PyMuPDF library converts PDF pages to high-resolution PNG images. Rendering pages is necessary since the original BEE guides consist of scanned, non-searchable document pages.", bold_prefix="1. Page Extraction & Rendering: ")
    add_bullet_styled(doc, "PyTesseract executes optical character recognition. The raw output text is indexed page-by-page. To prevent repeating the expensive OCR process on script reruns, each page's text is cached as a JSON record.", bold_prefix="2. OCR Analysis & Caching: ")
    add_bullet_styled(doc, "The system passes the text from the first 22 pages of each PDF (where the TOC is located) to Mistral Large. It prompts the model to return a structured JSON mapping containing the hierarchical Table of Contents, including exact starting and ending page numbers for every subchapter.", bold_prefix="3. TOC Structure Generation: ")
    add_bullet_styled(doc, "Using the page ranges extracted from the TOC, the system segments the complete OCR text of the book into individual section chunks. This creates logical boundary chunking (e.g., 'Chapter 3: Section 2 - Compressed Air Systems') rather than arbitrary word-count chunking.", bold_prefix="4. TOC-Aware Chunking: ")
    add_bullet_styled(doc, "Mistral LLM reads the raw text of each subsection chunk and generates a 2-3 sentence overview. This overview captures the semantic core of the subsection.", bold_prefix="5. Semantic Summarization: ")
    add_bullet_styled(doc, "The hierarchical node metadata and summaries are written to `tree.json` (under 100KB, loaded fully in memory). The raw page text is written line-by-line to a JSON Lines file `pages.jsonl`.", bold_prefix="6. Serialization: ")

    add_heading_styled(doc, "4.2 Phase 2: Runtime Query Execution Process", level=2)
    p_proc2 = doc.add_paragraph()
    p_proc2.paragraph_format.space_after = Pt(10)
    run_pp2 = p_proc2.add_run(
        "At runtime, the chatbot processes user questions in a linear pipeline:"
    )
    set_run_font(run_pp2, "Raleway")
    
    add_bullet_styled(doc, "The server receives the user's question through a REST POST call to `/api/chat`.", bold_prefix="1. API Endpoint: ")
    add_bullet_styled(doc, "The query is analyzed against the `tree.json` hierarchy. Mistral Large is prompted to inspect the chapter titles and subsection summaries and return a list of 1 to 3 subsection IDs that might contain the answer. This is the 'Vectorless' retrieval phase: it uses semantic LLM reasoning instead of cosine similarity.", bold_prefix="2. TOC Routing: ")
    add_bullet_styled(doc, "The system reads `pages.jsonl` and loads the exact raw OCR text corresponding to the chosen subsection IDs. This text becomes the 'Context'.", bold_prefix="3. Context Assembly: ")
    add_bullet_styled(doc, "The query and context are checked. If the list of selected sections is empty, the pipeline stops and responds with a pre-configured 'out-of-bounds' refutation.", bold_prefix="4. Pre-Generation Guardrail: ")
    add_bullet_styled(doc, "The query and contexts are formatted into the grounding prompt. Mistral Large generates a response and is instructed to return: (1) the grounded answer, (2) the list of source node IDs used, and (3) a boolean flag verifying that the answer is completely sourced from the context.", bold_prefix="5. Response Generation: ")
    add_bullet_styled(doc, "The post-generation module inspects the LLM response. It confirms that: (1) the grounding boolean is true, (2) all returned citation IDs exist in the retrieved context nodes, and (3) the text contains no hallucinated page numbers or figures. If any check fails, the answer is replaced with the out-of-bounds disclaimer.", bold_prefix="6. Post-Generation Validation: ")
    add_bullet_styled(doc, "The FastAPI backend appends specific metadata (such as exact book title, chapter title, and page range) to the citations and returns the complete JSON payload.", bold_prefix="7. Payload Construction: ")

    # ------------------ DESIGN DECISIONS ------------------
    add_heading_styled(doc, "5. Core Design Decisions", level=1)
    
    p_dd = doc.add_paragraph()
    run_dd = p_dd.add_run(
        "The development of Carbon Tatva chatbot was guided by specific design constraints regarding performance, hosting, and data accuracy:"
    )
    set_run_font(run_dd, "Raleway")

    # Table: Design Decisions
    table3 = doc.add_table(rows=7, cols=2)
    table3.autofit = False
    
    widths3 = [Inches(2.0), Inches(5.0)]
    for row in table3.rows:
        for i, width in enumerate(widths3):
            row.cells[i].width = width

    headers3 = ["Design Decision", "Rationale & Technical Justification"]
    for i, h in enumerate(headers3):
        cell_run = table3.cell(0, i).paragraphs[0].add_run(h)
        set_run_font(cell_run, "Raleway", size_pt=10, color_rgb=(255, 255, 255), bold=True)

    dd_data = [
        ["Elimination of Vector DB", "Using a vector database introduces database dependencies, connection pooling overhead, embedding model synchronization, and loss of document structure. TOC tree retrieval preserves structural context and uses Mistral's reasoning capabilities directly."],
        ["Pre-built Ingest Index", "FastAPI server runs on Vercel Serverless. Vercel functions have strict execution timeout limits and lack local Tesseract OCR binaries. By running the ingestion process offline and committing the completed `tree.json` and `pages.jsonl` files directly to the repository, Vercel only handles runtime routing and generation, executing in milliseconds."],
        ["OCR Text Caching", "Running PyTesseract OCR on 500+ pages of dense guidebooks takes over an hour. Implementing a cache file per page enables rapid iterative testing on the ingestion scripts without reprocessing pages."],
        ["Dual Guardrail Enforcement", "Technical guidebooks contain exact specifications, calculations, and safety values. Hallucinating a formula or standard could be dangerous. The dual-stage guardrail prevents out-of-domain answers and ensures citation safety."],
        ["Mistral Client Retries", "Mistral AI API can suffer from occasional rate limiting (HTTP 429) or temporary server errors. The `mistral_client.py` contains a custom retry wrapper with exponential backoff and jitter to ensure robust service."],
        ["Raleway Font & Carbon Tatva UI Theme", "To match the Carbon Tatva brand identity, the client interface is styled with a custom Navy/Teal palette. Raleway (a geometric sans-serif) is loaded and set as the sole typeface across all layouts to guarantee consistent visual branding."]
    ]

    for r_idx, row_vals in enumerate(dd_data):
        for c_idx, val in enumerate(row_vals):
            p = table3.cell(r_idx + 1, c_idx).paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            cell_run = p.add_run(val)
            set_run_font(cell_run, "Raleway", size_pt=9.5, color_rgb=(15, 25, 35), bold=(c_idx == 0))
            
    apply_table_styles(table3)

    # ------------------ REQUIREMENTS ------------------
    add_heading_styled(doc, "6. System Requirements & Dependencies", level=1)
    
    p_req = doc.add_paragraph()
    run_req = p_req.add_run(
        "To build, run, or re-ingest data for the Carbon Tatva Vectorless RAG chatbot, the following libraries and system tools are required:"
    )
    set_run_font(run_req, "Raleway")
    
    add_bullet_styled(doc, "Python 3.12 or higher (core language environment)", bold_prefix="Language Environment: ")
    add_bullet_styled(doc, "FastAPI & Uvicorn (web server routing and ASGI server interface)", bold_prefix="Web Framework: ")
    add_bullet_styled(doc, "PyMuPDF & pdf2image (handling PDF conversion and rendering)", bold_prefix="PDF Processing: ")
    add_bullet_styled(doc, "PyTesseract & Tesseract OCR engine (external binary required on system path for OCR extraction)", bold_prefix="OCR Engine: ")
    add_bullet_styled(doc, "Pillow (image manipulation and prep for Tesseract)", bold_prefix="Image Library: ")
    add_bullet_styled(doc, "Pydantic & Pydantic-Settings (settings management and strict type validation for environments and models)", bold_prefix="Configuration: ")
    add_bullet_styled(doc, "Python-docx (used to compile and build formal architecture documentation)", bold_prefix="Documentation: ")

    p_sys = doc.add_paragraph()
    p_sys.paragraph_format.space_before = Pt(10)
    run_sys = p_sys.add_run(
        "Important Note on System Binaries: Re-running the ingestion scripts requires having the Tesseract OCR engine installed and Poppler (for pdf2image) "
        "configured in the OS environment PATH. Running the server and querying the pre-built indexes does NOT require these system binaries."
    )
    set_run_font(run_sys, "Raleway", size_pt=10, color_rgb=(15, 25, 35), italic=True)

    # Save Document
    filename = "Carbon_Tatva_Architecture_Document.docx"
    doc.save(filename)
    print(f"Document saved successfully as: {filename}")

if __name__ == "__main__":
    main()
